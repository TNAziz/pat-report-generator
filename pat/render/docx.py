"""DOCX renderer.

Builds a Word document directly from the `Report` IR using
`python-docx`. We avoid the Markdown -> DOCX conversion path so the
output remains styled and editable (real tables, real heading levels,
real bold runs) rather than a wall of mono-spaced text.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from . import markdown as M
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .model import (
    Chart,
    Heatmap,
    NamedTable,
    NarrativeBlock,
    ProgramSection,
    Report,
    SemesterSection,
)

_RED = RGBColor(0xA0, 0x00, 0x00)

# XML 1.0 forbids most C0 control characters. Comments pasted from PDFs
# or older Word documents sometimes contain NUL, VT, or FF — save() then
# raises deep inside python-docx. Tab (\t), LF (\n), and CR (\r) stay in.
_XML_ILLEGAL_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _safe_text(s) -> str:
    """Strip XML-illegal control chars so python-docx's save() can't crash on them."""
    if s is None:
        return ""
    return _XML_ILLEGAL_CTRL.sub("", str(s))


def _add_percent_run(paragraph, val, below_threshold: bool = False):
    """Append a percent cell to a paragraph, bolded + red if below threshold."""
    run = paragraph.add_run(M._fmt_percent(val))
    if below_threshold:
        run.bold = True
        run.font.color.rgb = _RED
    return run


def _add_summary_table(doc, section: ProgramSection):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, name in enumerate(
        ["Semester", "Sub-Outcome", "Performance Indicator", "Performance"]
    ):
        run = hdr[i].paragraphs[0].add_run(name)
        run.bold = True
    if not section.summary:
        row = table.add_row().cells
        row[0].text = "No data for this program."
        return
    for r in section.summary:
        cells = table.add_row().cells
        cells[0].text = r.semester
        cells[1].text = r.suboutcome
        _add_percent_run(cells[2].paragraphs[0], r.performance_indicator, r.below_threshold)
        _add_percent_run(cells[3].paragraphs[0], r.performance, r.below_threshold)
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "Cells in bold red indicate the average performance was below "
        "the performance indicator."
    )
    foot_run.italic = True
    foot_run.font.size = Pt(9)
    foot_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_measure(doc, m):
    doc.add_heading(f"Sub-Outcome: {m.suboutcome}", level=4)

    p = doc.add_paragraph()
    p.add_run("Measure Description: ").bold = True
    p.add_run(_safe_text(m.measure_description) or "N/A")

    p = doc.add_paragraph()
    p.add_run("Performance Threshold: ").bold = True
    _add_percent_run(p, m.performance_indicator, m.below_threshold)

    p = doc.add_paragraph()
    p.add_run("Student Performance: ").bold = True
    _add_percent_run(p, m.performance, m.below_threshold)

    p = doc.add_paragraph()
    p.add_run("n = ").bold = True
    p.add_run(str(m.n) if m.n is not None else "N/A")

    quote = doc.add_paragraph(style="Intense Quote")
    quote.add_run("Comments: ").bold = True
    quote.add_run(_safe_text(m.comments))
    quote.add_run("\n")
    quote.add_run("Actions Taken: ").bold = True
    quote.add_run(_safe_text(m.actions_taken))


def _add_semester_section(doc, section: SemesterSection):
    doc.add_heading(section.semester, level=2)
    p = doc.add_paragraph()
    p.add_run("Instructor: ").bold = True
    p.add_run(section.instructor or "N/A")
    for m in section.measures:
        _add_measure(doc, m)


def _add_program_section(doc, section: ProgramSection):
    doc.add_heading(section.program_label, level=1)
    _add_summary_table(doc, section)
    for sem in section.semesters:
        _add_semester_section(doc, sem)


def _add_named_table(doc, t: NamedTable):
    if t.title:
        doc.add_heading(t.title, level=2)
    if not t.columns:
        return
    table = doc.add_table(rows=1, cols=len(t.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(t.columns):
        run = hdr[i].paragraphs[0].add_run(col)
        run.bold = True
    for row in t.rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = _safe_text(val)
    if t.footnote:
        p = doc.add_paragraph()
        run = p.add_run(t.footnote)
        run.italic = True
        run.font.size = Pt(9)


_MD_LIST_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_paragraph(doc, md: str):
    """Add a paragraph of markdown to the doc.

    Supports a small markdown subset: bullet lists (- foo) and inline
    **bold**. The narrative blocks in the spec only need this much; if
    we ever need richer markdown, swap to the markdown library +
    html-to-docx, but that adds dependencies for one feature.
    """
    for line in md.split("\n"):
        m = _MD_LIST_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(1))
        elif line.strip():
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
        else:
            doc.add_paragraph()


def _add_inline_runs(paragraph, text: str):
    """Add text with **bold** markdown converted to bold runs."""
    pos = 0
    for m in _MD_BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_narrative(doc, block: NarrativeBlock):
    if block.heading:
        level = max(1, min(9, getattr(block, "level", 2)))
        doc.add_heading(block.heading, level=level)
    _add_markdown_paragraph(doc, block.body_markdown)


def _add_body_item(doc, item):
    """Add one ordered ``Report.body`` item by type."""
    if isinstance(item, NarrativeBlock):
        _add_narrative(doc, item)
    elif isinstance(item, NamedTable):
        _add_named_table(doc, item)
    elif isinstance(item, Chart):
        _add_chart(doc, item)
    elif isinstance(item, Heatmap):
        _add_heatmap(doc, item)
    else:
        raise TypeError(f"unsupported Report.body item: {type(item).__name__}")


def _add_chart(doc, chart: Chart):
    """Render chart as a simple data table in DOCX.

    Word-native charting is brittle and varies across Word versions;
    embedding the chart as a table preserves the underlying data and
    looks clean. If we later want a real chart image, swap this to a
    matplotlib PNG insert -- the renderer signature stays the same.
    """
    if not chart.series:
        return
    doc.add_heading(chart.title, level=2)
    all_x = sorted({x for s in chart.series for x in s.x})
    table = doc.add_table(rows=1, cols=1 + len(chart.series))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run(chart.x_label).bold = True
    for i, s in enumerate(chart.series, start=1):
        hdr[i].paragraphs[0].add_run(s.name).bold = True
    for x in all_x:
        cells = table.add_row().cells
        cells[0].text = str(x)
        for i, s in enumerate(chart.series, start=1):
            try:
                idx = s.x.index(x)
                cells[i].text = str(s.y[idx])
            except ValueError:
                cells[i].text = ""




_RAMPS = {
    "blues":  ((247, 251, 255), (8,  48,  107)),
    "greens": ((247, 252, 245), (0,  68,  27)),
    "reds":   ((255, 245, 240), (103, 0,   13)),
}


def _shade_cell(cell, hex_color: str):
    """Set the background fill color of a docx table cell.

    python-docx doesn't expose this directly; we drop down to the raw
    OOXML element for the cell's properties.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _rgb_to_hex(rgb):
    return "{:02X}{:02X}{:02X}".format(*rgb)


def _interp(low, high, t):
    return tuple(int(low[i] + (high[i] - low[i]) * t) for i in range(3))


def _add_heatmap(doc, h: Heatmap):
    """Render a Heatmap as a real Word table with cell shading."""
    if not h.rows or not h.columns:
        return
    doc.add_heading(h.title, level=2)
    flat = [v for row in h.values for v in row
            if v is not None and not (isinstance(v, float) and v != v)]
    vmin = h.vmin if h.vmin is not None else (min(flat) if flat else 0)
    vmax = h.vmax if h.vmax is not None else (max(flat) if flat else 1)
    if vmax <= vmin:
        vmax = vmin + 1
    low_rgb, high_rgb = _RAMPS.get(h.color_scheme, _RAMPS["blues"])

    table = doc.add_table(rows=1, cols=1 + len(h.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run(h.row_label or "").bold = True
    for i, c in enumerate(h.columns, start=1):
        hdr[i].paragraphs[0].add_run(str(c)).bold = True

    for r, label in enumerate(h.rows):
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(str(label)).bold = True
        for c in range(len(h.columns)):
            try:
                v = h.values[r][c]
            except IndexError:
                v = None
            is_zero = (v == 0)
            is_missing = (v is None or (isinstance(v, float) and v != v))
            if is_missing:
                display = h.empty_marker
            else:
                display = h.value_format.format(v)
            if h.highlight_zero and is_zero:
                # Pink background + bold red text -> "no coverage".
                _shade_cell(cells[c + 1], "FCE6E6")
                run = cells[c + 1].paragraphs[0].add_run(display)
                run.bold = True
                run.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
            else:
                if is_missing:
                    t = 0.0
                else:
                    t = max(0.0, min(1.0, (v - vmin) / (vmax - vmin)))
                bg = _interp(low_rgb, high_rgb, t)
                _shade_cell(cells[c + 1], _rgb_to_hex(bg))
                lum = 0.299 * bg[0] + 0.587 * bg[1] + 0.114 * bg[2]
                run = cells[c + 1].paragraphs[0].add_run(display)
                if lum <= 130:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    if h.caption:
        p = doc.add_paragraph()
        run = p.add_run(h.caption)
        run.italic = True
        run.font.size = Pt(9)


def render(report: Report) -> bytes:
    """Render a Report to a Word document, returned as bytes."""
    doc = Document()
    # Tighten the default margins slightly so tables don't overflow.
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    title = doc.add_heading(report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if report.subtitle:
        p = doc.add_paragraph()
        run = p.add_run(report.subtitle)
        run.italic = True
        run.font.size = Pt(12)
    if report.generated_on:
        p = doc.add_paragraph()
        run = p.add_run(f"Generated {report.generated_on.strftime('%B %d, %Y')}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    for section in report.sections:
        _add_program_section(doc, section)
    for item in report.body:
        _add_body_item(doc, item)
    for block in report.narrative:
        _add_narrative(doc, block)
    for table in report.tables:
        _add_named_table(doc, table)
    for chart in report.charts:
        _add_chart(doc, chart)
    for heatmap in report.heatmaps:
        _add_heatmap(doc, heatmap)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
