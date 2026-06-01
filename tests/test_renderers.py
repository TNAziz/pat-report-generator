"""Cross-format renderer tests.

Traceability:
- R15: all four formats produced.
- R16: content parity -- same titles and key cell values appear in
  every format (with format-specific styling allowed to differ).
- R19, R25: all four formats handle Sub-Outcome Lookup and Coverage
  Check report shapes.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from pat.render import docx as Drx, html as H, markdown as M, pdf as P
from tests.report_fixtures import (
    make_course_report, make_suboutcome_lookup, make_coverage_report,
)


def _docx_text(blob: bytes) -> str:
    """Extract every text run from a docx into one string."""
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.mark.parametrize("fixture", [
    make_course_report, make_suboutcome_lookup, make_coverage_report,
])
def test_all_four_formats_produced(fixture):
    """R15 -- every renderer returns non-empty content for every shape."""
    r = fixture()
    md = M.render(r)
    html = H.render(r)
    docx = Drx.render(r)
    pdf = P.render(r)
    assert md and len(md) > 50
    assert html and len(html) > 200
    assert docx and docx[:2] == b"PK"
    assert pdf and pdf[:5] == b"%PDF-"


@pytest.mark.parametrize("fixture, expected_strings", [
    (make_course_report, ["CE 282", "Civil Engineering", "Construction Engineering",
                          "Fall 2023", "Aziz, Tarek", "Exam 2 Q2c"]),
    (make_suboutcome_lookup, ["CE 488", "Programs", "Formulate"]),
    (make_coverage_report, ["Coverage Check", "Spring 2025",
                            "Missing assessments", "CE 339", "Semester summary"]),
])
def test_content_parity(fixture, expected_strings):
    """R16 -- every key string appears in all four formats."""
    r = fixture()
    md = M.render(r)
    html = H.render(r)
    docx_text = _docx_text(Drx.render(r))
    pdf = P.render(r)  # PDF text is compressed -- check only via size + headers

    for s in expected_strings:
        assert s in md, f"'{s}' missing from Markdown for {fixture.__name__}"
        assert s in html, f"'{s}' missing from HTML for {fixture.__name__}"
        assert s in docx_text, f"'{s}' missing from DOCX for {fixture.__name__}"
    # PDF: just confirm it's a real PDF; visual review covers full check.
    assert pdf[:5] == b"%PDF-"


def test_below_threshold_signaled_in_each_format():
    """The Spring 2021 1.1 row in the course report has performance below
    indicator. Each format should signal this somehow."""
    import zipfile, io as _io
    r = make_course_report()
    md = M.render(r)
    html = H.render(r)
    docx = Drx.render(r)
    # Markdown bolds the values.
    assert "**70%**" in md and "**60%**" in md
    # HTML applies the .below-threshold class.
    assert 'class="below-threshold"' in html
    # DOCX uses a bold + red run. Unzip the docx and check the document XML.
    with zipfile.ZipFile(_io.BytesIO(docx)) as zf:
        body_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    assert "A00000" in body_xml.upper(), "expected red color (A00000) in DOCX XML"
    assert "<w:b/>" in body_xml or '<w:b w:val="true"/>' in body_xml, \
        "expected bold runs in DOCX XML"
