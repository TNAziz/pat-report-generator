"""Unit tests for pat.render.docx."""

from __future__ import annotations

import io

from docx import Document

from pat.render import docx as D
from pat.render.model import (
    MeasureDetail, NamedTable, ProgramSection, Report, SemesterSection,
)
from tests.report_fixtures import (
    make_course_report, make_suboutcome_lookup, make_coverage_report,
)


def _open(blob: bytes):
    return Document(io.BytesIO(blob))


def test_render_returns_valid_docx_bytes():
    blob = D.render(make_course_report())
    # DOCX files are ZIP archives starting with 'PK'.
    assert blob[:2] == b"PK"
    doc = _open(blob)
    assert doc.paragraphs  # at least the title


def test_render_includes_title_and_subtitle():
    doc = _open(D.render(make_course_report()))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "CE 282" in text
    assert "Spring 2020 – Spring 2026" in text


def test_render_has_real_tables():
    doc = _open(D.render(make_course_report()))
    # The CE summary table + the CON summary table + ... at least 2.
    assert len(doc.tables) >= 2
    # Headers in the first summary table.
    first = doc.tables[0]
    headers = [c.text for c in first.rows[0].cells]
    assert headers == ["Semester", "Sub-Outcome", "Performance Indicator", "Performance"]


def test_render_marks_below_threshold_with_bold_red():
    doc = _open(D.render(make_course_report()))
    # Walk every run in every cell, look for a bold red percent.
    found = False
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold and run.font.color and run.font.color.rgb is not None:
                            if "60%" in run.text or "70%" in run.text:
                                found = True
                                break
    assert found, "Expected at least one bold+red percent run for below-threshold rows"


def test_render_suboutcome_lookup_has_bullet_list():
    doc = _open(D.render(make_suboutcome_lookup()))
    # python-docx exposes list-style paragraphs via style name.
    bullet_paras = [p for p in doc.paragraphs if p.style.name.startswith("List")]
    assert bullet_paras, "Expected at least one bulleted paragraph"


def test_render_coverage_includes_chart_data_table():
    doc = _open(D.render(make_coverage_report()))
    # Expect 3 tables: missing, summary, chart-data.
    assert len(doc.tables) == 3
    chart_table = doc.tables[-1]
    headers = [c.text for c in chart_table.rows[0].cells]
    assert headers[0] == "Year"
    assert "CE" in headers and "CON" in headers and "ENE" in headers


def test_render_empty_report_does_not_crash():
    blob = D.render(Report(title="Empty"))
    doc = _open(blob)
    assert any("Empty" in p.text for p in doc.paragraphs)


# -------- control-char sanitization --------


def _report_with_comment(text: str) -> Report:
    return Report(
        title="Ctrl-char probe",
        sections=[
            ProgramSection(
                program_code="CE",
                program_label="Civil Engineering",
                summary=[],
                semesters=[
                    SemesterSection(
                        semester="Fall 2024",
                        instructor="Aziz, Tarek",
                        measures=[
                            MeasureDetail(
                                suboutcome="1.1",
                                measure_description="Q1",
                                performance_indicator=70,
                                performance=80,
                                n=10,
                                comments=text,
                                actions_taken="ok",
                            )
                        ],
                    )
                ],
            )
        ],
    )


def test_render_survives_nul_and_vt_in_comments():
    """A comment pasted from a PDF or older Word doc can contain \\x00 / \\x0b /
    \\x0c; python-docx's save() raises deep inside if we let those through.
    """
    blob = D.render(_report_with_comment("bad \x00 chars \x0b here \x0c end"))
    assert blob[:2] == b"PK"
    doc = _open(blob)
    text = "\n".join(p.text for c in [cell for tbl in doc.tables for row in tbl.rows for cell in row.cells] for p in c.paragraphs) + \
        "\n".join(p.text for p in doc.paragraphs)
    assert "bad  chars  here  end" in text


def test_named_table_cell_control_chars_are_stripped():
    tbl = NamedTable(
        title="T",
        columns=["A"],
        rows=[["x\x00y"]],
        footnote=None,
    )
    blob = D.render(Report(title="T", tables=[tbl]))
    doc = _open(blob)
    cell_texts = [c.text for tb in doc.tables for row in tb.rows for c in row.cells]
    assert "xy" in cell_texts
